#!/usr/bin/env python3
"""Render the approved Navy SRCLA Markdown proposal as a polished DOCX."""

from __future__ import annotations

import re
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


TITLE = "Navy – Ví blockchain tích hợp farming USDC tối ưu bằng thuật toán SRCLA trên Base"


@dataclass(frozen=True)
class ContentBlock:
    kind: str
    text: str


@dataclass(frozen=True)
class ScheduleRow:
    stage: str
    period: str
    work: str
    result: str


@dataclass(frozen=True)
class ProposalContent:
    title: str
    metadata: dict[str, str]
    blocks: list[ContentBlock]
    schedule: list[ScheduleRow]
    signatures: dict[str, str]


def _strip_markdown(value: str) -> str:
    return re.sub(r"\*\*([^*]+)\*\*|`([^`]+)`", lambda m: m.group(1) or m.group(2), value)


def _parse_label(line: str) -> tuple[str, str] | None:
    match = re.fullmatch(r"\*\*([^*]+):\*\*\s*(.+?)\s{0,2}", line)
    if not match:
        return None
    return match.group(1), match.group(2).rstrip()


def _parse_table_row(line: str) -> list[str] | None:
    if not line.startswith("|"):
        return None
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    if len(cells) != 4 or all(set(cell) <= {"-", ":"} for cell in cells):
        return None
    return cells


def parse_source(path: Path) -> ProposalContent:
    lines = path.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("source must begin with a level-one title")
    title = lines[0][2:].strip()
    if title != TITLE:
        raise ValueError(f"unexpected title: {title}")

    metadata: dict[str, str] = {}
    signatures: dict[str, str] = {}
    blocks: list[ContentBlock] = []
    schedule: list[ScheduleRow] = []
    section_number = 0
    for raw_line in lines[1:]:
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            number_match = re.match(r"(\d+)\.", heading)
            section_number = int(number_match.group(1)) if number_match else 0
            if section_number <= 7:
                blocks.append(ContentBlock("section", heading))
            continue
        if section_number == 8:
            cells = _parse_table_row(line)
            if cells and cells != ["Giai đoạn", "Thời gian", "Công việc", "Kết quả"]:
                schedule.append(ScheduleRow(*cells))
            continue
        if section_number == 9:
            parsed = _parse_label(line)
            if parsed:
                signatures[parsed[0]] = parsed[1]
            continue
        if section_number == 0:
            parsed = _parse_label(line)
            if parsed:
                metadata[parsed[0]] = parsed[1]
            continue
        if line.startswith("### "):
            blocks.append(ContentBlock("subsection", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(ContentBlock("bullet", line[2:].strip()))
        else:
            blocks.append(ContentBlock("paragraph", line))

    if not schedule:
        raise ValueError("source contains no schedule rows")
    return ProposalContent(title, metadata, blocks, schedule, signatures)


def _set_run_font(run, name: str = "Times New Roman", size: int = 13) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in re.finditer(r"\*\*([^*]+)\*\*|`([^`]+)`", text):
        if match.start() > cursor:
            _set_run_font(paragraph.add_run(text[cursor : match.start()]))
        if match.group(1) is not None:
            run = paragraph.add_run(match.group(1))
            run.bold = True
            _set_run_font(run)
        else:
            run = paragraph.add_run(match.group(2))
            _set_run_font(run, "Courier New", 11)
        cursor = match.end()
    if cursor < len(text):
        _set_run_font(paragraph.add_run(text[cursor:]))


def _set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: int = 12) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(_strip_markdown(text))
    run.bold = bold
    _set_run_font(run, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _remove_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    properties.append(cant_split)


def configure_page(document: DocumentType) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)


def configure_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for name, size, before, after in [
        ("Proposal Heading 1", 14, 10, 4),
        ("Proposal Heading 2", 13, 6, 3),
    ]:
        if name in document.styles:
            style = document.styles[name]
        else:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    list_style = document.styles["List Bullet"]
    list_style.font.name = "Times New Roman"
    list_style.font.size = Pt(13)
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    list_style.paragraph_format.left_indent = Cm(0.7)
    list_style.paragraph_format.first_line_indent = Cm(-0.4)
    list_style.paragraph_format.space_after = Pt(2)


def add_administrative_header(document: DocumentType, content: ProposalContent) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(8.2)
    table.columns[1].width = Cm(8.2)
    _remove_table_borders(table)

    left = table.cell(0, 0)
    left.text = ""
    for index, line in enumerate(
        ["ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH", "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN"]
    ):
        paragraph = left.paragraphs[0] if index == 0 else left.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = True
        _set_run_font(run, size=11)

    right = table.cell(0, 1)
    right.text = ""
    for index, (line, bold) in enumerate(
        [
            ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", True),
            ("Độc lập – Tự do – Hạnh phúc", True),
        ]
    ):
        paragraph = right.paragraphs[0] if index == 0 else right.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        _set_run_font(run, size=11)


def add_title_and_metadata(document: DocumentType, content: ProposalContent) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run("ĐỀ CƯƠNG CHI TIẾT")
    run.bold = True
    _set_run_font(run, size=18)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(content.title)
    run.bold = True
    _set_run_font(run, size=15)

    table = document.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.2)
    _remove_table_borders(table)
    values = [
        ("Cán bộ hướng dẫn", content.metadata["Cán bộ hướng dẫn"]),
        ("Thời gian thực hiện", content.metadata["Thời gian thực hiện"]),
        ("Sinh viên thực hiện", content.metadata["Sinh viên thực hiện"].replace("; ", "\n")),
    ]
    for row, (label, value) in zip(table.rows, values):
        _set_cell_text(row.cells[0], f"{label}:", bold=True, size=13)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[1].text = ""
        for index, line in enumerate(value.splitlines()):
            paragraph = row.cells[1].paragraphs[0] if index == 0 else row.cells[1].add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(line)
            _set_run_font(run, size=13)

    body_title = document.add_paragraph()
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_title.paragraph_format.space_before = Pt(10)
    body_title.paragraph_format.space_after = Pt(6)
    body_title.paragraph_format.keep_with_next = True
    run = body_title.add_run("NỘI DUNG ĐỀ TÀI")
    run.bold = True
    _set_run_font(run, size=14)


def add_academic_body(document: DocumentType, content: ProposalContent) -> None:
    for block in content.blocks:
        if block.kind == "section":
            paragraph = document.add_paragraph(style="Proposal Heading 1")
            _add_inline(paragraph, block.text)
        elif block.kind == "subsection":
            paragraph = document.add_paragraph(style="Proposal Heading 2")
            _add_inline(paragraph, block.text)
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, block.text)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(1)
            _add_inline(paragraph, block.text)


def add_schedule(document: DocumentType, rows: list[ScheduleRow]) -> None:
    document.add_page_break()
    heading = document.add_paragraph(style="Proposal Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    heading.add_run("8. Kế hoạch thực hiện")

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.0), Cm(3.0), Cm(6.3), Cm(4.3)]
    for column, width in zip(table.columns, widths):
        column.width = width
    headers = ["Giai đoạn", "Thời gian", "Công việc", "Kết quả"]
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, bold=True, center=True, size=11)
    _repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])

    for item in rows:
        row = table.add_row()
        values = [item.stage, item.period, item.work, item.result]
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            _set_cell_text(cell, value, center=index == 1, size=10)
        _prevent_row_split(row)


def add_signature_block(document: DocumentType, content: ProposalContent) -> None:
    heading = document.add_paragraph(style="Proposal Heading 1")
    heading.paragraph_format.space_before = Pt(12)
    heading.add_run("XÁC NHẬN")

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date.add_run(content.signatures.get("Địa điểm, ngày lập đề cương", ""))

    table = document.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for column, width in zip(table.columns, [Cm(5.4), Cm(5.4), Cm(5.4)]):
        column.width = width
    _remove_table_borders(table)
    labels = ["Xác nhận của CBHD", "Sinh viên 1", "Sinh viên 2"]
    names = [
        content.signatures.get("Cán bộ hướng dẫn", ""),
        content.signatures.get("Sinh viên 1", ""),
        content.signatures.get("Sinh viên 2", ""),
    ]
    for index, label in enumerate(labels):
        _set_cell_text(table.cell(0, index), label, bold=True, center=True, size=12)
        _set_cell_text(table.cell(1, index), "(Ký tên và ghi rõ họ tên)", center=True, size=11)
        table.cell(1, index).paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        table.cell(1, index).paragraphs[0].add_run().add_break(WD_BREAK.LINE)
        _set_cell_text(table.cell(2, index), names[index], bold=True, center=True, size=12)


def build(source_md: Path, template_docx: Path, output_docx: Path) -> None:
    if not template_docx.exists():
        raise FileNotFoundError(f"template does not exist: {template_docx}")
    if template_docx.resolve() == output_docx.resolve():
        raise ValueError("output path must not overwrite the source template")
    Document(template_docx)  # Validate that the supplied administrative reference can be opened.
    content = parse_source(source_md)

    document = Document()
    configure_page(document)
    configure_styles(document)
    document.core_properties.title = content.title
    document.core_properties.subject = "Đề cương chi tiết đồ án Navy SRCLA"
    document.core_properties.author = "Nguyễn Ngọc Anh Khoa; Trương Nguyễn Thùy Anh"
    add_administrative_header(document, content)
    add_title_and_metadata(document, content)
    add_academic_body(document, content)
    add_schedule(document, content.schedule)
    add_signature_block(document, content)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile(
        prefix=f".{output_docx.stem}-", suffix=".docx", dir=output_docx.parent, delete=False
    ) as temporary:
        temporary_path = Path(temporary.name)
    try:
        document.save(temporary_path)
        Document(temporary_path)
        temporary_path.replace(output_docx)
    finally:
        temporary_path.unlink(missing_ok=True)


def main(argv: list[str]) -> int:
    if len(argv) != 4:
        print("usage: build_docx.py SOURCE_MD TEMPLATE_DOCX OUTPUT_DOCX", file=sys.stderr)
        return 2
    source, template, output = map(Path, argv[1:])
    build(source, template, output)
    print(f"WROTE {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
