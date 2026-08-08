#!/usr/bin/env python3
"""Integration tests for the Navy SRCLA proposal DOCX renderer."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
BUILDER = ROOT / "tools/docx_proposal/build_docx.py"
SOURCE = ROOT / "docs/reports/navy-srcla-detailed-proposal.md"
TEMPLATE = ROOT / "De-cuong-chi-tiet.docx"
EXACT_TITLE = "Navy – Ví blockchain tích hợp farming USDC tối ưu bằng thuật toán SRCLA trên Base"


def document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class BuildDocxTest(unittest.TestCase):
    def test_build_creates_valid_administrative_document(self) -> None:
        """Catches missing/invalid output, wrong form structure, or lost approved copy."""
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "proposal.docx"
            result = subprocess.run(
                [sys.executable, str(BUILDER), str(SOURCE), str(TEMPLATE), str(output)],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertTrue(output.exists(), "builder returned success without writing output")

            document = Document(output)
            self.assertAlmostEqual(document.sections[0].page_height.mm, 297.0, places=0)
            self.assertAlmostEqual(document.sections[0].page_width.mm, 210.0, places=0)
            self.assertGreaterEqual(len(document.tables), 3)

            all_text = document_text(document)
            for required in [
                EXACT_TITLE,
                "ĐỀ CƯƠNG CHI TIẾT",
                "ThS. Nguyễn Tấn Toàn",
                "Nguyễn Ngọc Anh Khoa",
                "Trương Nguyễn Thùy Anh",
                "Safe, Robust, Cost-Aware Lending Allocator",
                "Base Sepolia",
                "Base mainnet fork",
            ]:
                self.assertIn(required, all_text)

            headings = [
                "1. Lý do chọn đề tài",
                "2. Mục tiêu",
                "3. Phạm vi và đối tượng sử dụng",
                "4. Phương pháp thực hiện",
                "5. Nền tảng công nghệ",
                "6. Kết quả mong đợi",
                "7. Hướng phát triển",
                "8. Kế hoạch thực hiện",
            ]
            positions = [all_text.index(heading) for heading in headings]
            self.assertEqual(positions, sorted(positions))

            schedule = next(
                table for table in document.tables if table.cell(0, 0).text == "Giai đoạn"
            )
            self.assertEqual(len(schedule.columns), 4)
            self.assertEqual(schedule.rows[0].cells[3].text, "Kết quả")
            self.assertIn("03/09/2026", schedule.rows[1].cells[1].text)
            self.assertIn("26/12/2026", schedule.rows[-1].cells[1].text)
            self.assertLess(all_text.index("8. Kế hoạch thực hiện"), all_text.index("XÁC NHẬN"))

            for forbidden in ["TBD", "TODO", "điền", "31/11", "Ethereum Sepolia", "MongoDB"]:
                self.assertNotIn(forbidden, all_text)


if __name__ == "__main__":
    unittest.main()
